from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from .models import Maqola
from .forms import MaqolaForm
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from apps.accounts.models import Profile


def home(request):
    return render(request, 'home.html')


def create_maqola(request):
    if request.method == 'POST':
        form = MaqolaForm(request.POST)
        if form.is_valid():
            maqola = form.save(commit=False)
            maqola.user = request.user  # shu foydalanuvchiga biriktiramiz
            maqola.save()
            return redirect('maqola:list')
    else:
        form = MaqolaForm()
    return render(request, 'maqola/create.html', {'form': form})



def list_maqola(request):
    maqolalar = Maqola.objects.filter(user=request.user).order_by('id')
    return render(request, 'maqola/list.html', {'maqolalar': maqolalar})


def update_maqola(request, pk):
    maqola = get_object_or_404(Maqola, pk=pk, user=request.user)
    form = MaqolaForm(request.POST or None, instance=maqola)
    if request.method == 'POST' and form.is_valid():
        form.save()
        return redirect('maqola:list')
    return render(request, 'maqola/update.html', {'form': form})


def delete_maqola(request, pk):
    maqola = get_object_or_404(Maqola, pk=pk, user=request.user)
    if request.method == 'POST':
        maqola.delete()
        return redirect('maqola:list')
    return render(request, 'maqola/delete.html', {'maqola': maqola})


def export_maqola_docx(request):
    user = request.user
    maqolalar = Maqola.objects.filter(user=user).order_by('id')
    profile = Profile.objects.get(user=user)

    doc = Document()

    # LANDSCAPE format va sahifa chetlarini sozlash
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Sarlavha
    heading_text = (
        "Muhammad al-Xorazmiy nomidagi Toshkent axborot texnologiyalari universiteti Kiberxavfsizlik fakulteti\n"
        f"{profile.fakultet_raqami} – “{profile.fakultet} (Axborot-kommunikatsiya texnologiyalari va servis)” "
        f"ta’lim yo‘nalishi {profile.guruh_raqami}-guruh talabasi {profile.talaba_fish}ning\n"
        "ILMIY ISHLARI RO‘YXATI"
    )
    heading = doc.add_paragraph(heading_text)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in heading.runs:
        run.bold = True
        run.font.size = Pt(12)

    # Jadval
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.autofit = False

    column_widths = [Inches(0.5), Inches(2.0), Inches(1.2), Inches(2.5), Inches(1.2), Inches(1.6)]

    headers = [
        "№",
        "Ilmiy ishning nomi",
        "Bosma, qo‘lyozma yoki elektron",
        "Jurnal, to‘plam (vil., nomer, betlar), nashriyot yoki mualliflik guvohnomasi nomeri",
        "Bosma taboq yoki betlar soni, mualliflik ishtiroki",
        "Hammualliflarning familiyalari, ismlari, otalarining ismlari"
    ]

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].width = column_widths[i]

    for idx, m in enumerate(maqolalar, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = m.title
        row_cells[2].text = m.format
        row_cells[3].text = m.journal_information
        row_cells[4].text = m.piece
        row_cells[5].text = m.authors

        for i, width in enumerate(column_widths):
            row_cells[i].width = width

    # Export to DOCX
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=ilmiy_ishlar.docx'
    doc.save(response)
    return response